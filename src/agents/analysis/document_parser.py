"""
Document parser for PDF and DOCX files
"""

import os
from typing import Optional
import PyPDF2
import pdfplumber
from docx import Document
import logging

logger = logging.getLogger(__name__)

class DocumentParser:
    """Parse PDF and DOCX files to extract text content"""
    
    @staticmethod
    def extract_text_from_file(file_path: str) -> Optional[str]:
        """
        Extract text from PDF or DOCX file
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content or None if failed
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                return DocumentParser._extract_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return DocumentParser._extract_from_docx(file_path)
            else:
                logger.error(f"Unsupported file type: {file_ext}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF using pdfplumber (better than PyPDF2)"""
        text = ""
        
        try:
            # Try pdfplumber first (better text extraction)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():
                return text.strip()
                
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}: {e}")
        
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"PyPDF2 also failed for {file_path}: {e}")
            raise
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    
    @staticmethod
    def validate_file(file_path: str) -> bool:
        """
        Validate if file exists and is supported format
        
        Args:
            file_path: Path to the document file
            
        Returns:
            True if file is valid and supported
        """
        if not os.path.exists(file_path):
            return False
        
        file_ext = os.path.splitext(file_path)[1].lower()
        return file_ext in ['.pdf', '.docx', '.doc']